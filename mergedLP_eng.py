import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.shared import Pt, Inches
from io import BytesIO

# --- 1. CONFIGURATION ---
st.set_page_config(page_title="MERGED Master Planner", layout="wide")
st.title("🎓 PEDATI SMART LESSON PLANNER")

# --- MAIN PAGE CONFIGURATION & USER API KEY BAR (AT THE VERY TOP) ---
user_api_key = st.text_input(
    "🔑 ENTER YOUR GEMINI API KEY:", 
    type="password", 
    help="Get your API key from Google AI Studio using your Gmail account."
)

# Helper function to dynamically check and load models based on the user's key
def get_working_model(api_key):
    try:
        genai.configure(api_key=api_key)
        for m in genai.list_models():
            if 'generateContent' in m.supported_generation_methods:
                return m.name
    except Exception as e:
        st.error(f"INVALID API KEY OR CONNECTION ERROR: {str(e)}")
        return None
    return "models/gemini-1.5-flash"  # Default fallback


# Process model assignment if the key is provided
selected_model_name = None
if user_api_key:
    selected_model_name = get_working_model(user_api_key)
    if selected_model_name:
        st.info(f"SYSTEM CONNECTED VIA YOUR API KEY. ACTIVE MODEL: {selected_model_name.upper()}")
else:
    st.warning("⚠️ PLEASE ENTER YOUR PERSONAL GEMINI API KEY ABOVE TO START.")


# --- 2. AI GENERATION ENGINE ---
def generate_pedati_plan(topic, syllabus, extra_context, api_key, model_name):
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel(model_name)
    
    prompt = f"""
    Topic: {topic}. Syllabus Code: {syllabus}. Context: {extra_context}.
    Generate a lesson plan in English. 
    NO Malay terms.
    
    CRITICAL TEXT FORMATTING RULES:
    1. DO NOT use double asterisks (**) anywhere in the output.
    2. Ensure every single section title and stage marker is in full CAPITAL LETTERS.
    3. Use numbers for lists where appropriate.
    
    Use these exact stage names:
    P [Prior Knowledge], E [Engage], D [Develop], A [Apply], T [Test], I [Improve].

    Structure with these exact markers for boxing:
    SECTION: LESSON OBJECTIVES
    [4 points]
    SECTION: LESSON OUTCOMES
    [4 points]
    SECTION: SUCCESS CRITERIA
    [4 points]
    SECTION: PREREQUISITE
    [1 point]
    SECTION: KEYWORDS
    [6 items]
    SECTION: HOTS
    [any 4 main domains in the Bloom's taxonomy]
    SECTION: DIGITAL CITIZENSHIP
    [4 points on the use of online resources like youtube channel or canva application or use of chromebooks or use of digital devices]

    SECTION: PEDATI FLOW GRID
    BLOCK_START: P: PREPARATION (LEARN)
    LECTURER:
    [Activity lists]
    STUDENTS:
    [Activity lists]
    BLOCK_END
    
    BLOCK_START: E: ENGAGE (EXPLORE)
    LECTURER:
    [Activity lists]
    STUDENTS:
    [Activity lists]
    BLOCK_END

    BLOCK_START: D.A: DELIVER AND APPLY
    LECTURER:
    [Activity lists]
    STUDENTS:
    [Activity lists]
    BLOCK_END

    BLOCK_START: T.I: TEST AND EVALUATE
    LECTURER:
    [Activity lists]
    STUDENTS:
    [Activity lists]
    BLOCK_END
    """
    try:
        response = model.generate_content(prompt)
        return response.text.replace("**", "")
    except Exception as e:
        return f"SYSTEM ERROR: {str(e)}"


# --- 3. WORD DOCUMENT EXPORT ENGINE ---
def create_word_export(topic, syllabus, text):
    doc = Document()
    
    # Global document layout formatting override configurations
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)  # Set 14 points font size
    
    p_format = style.paragraph_format
    p_format.line_spacing = 1.0  # Set 1 paragraph / single spacing
    p_format.space_after = Pt(12)  # Maintain clean paragraph block dynamic gap

    # Document Header Title - FULL CAPITAL LETTERS
    title_p = doc.add_paragraph()
    run_title = title_p.add_run(f'LESSON PLAN: {topic.upper()} ({syllabus.upper()})')
    run_title.bold = True
    run_title.font.size = Pt(16)

    # 1. Admin Header Table (6-field layout)
    admin_table = doc.add_table(rows=3, cols=4)
    admin_table.style = 'Table Grid'
    labels = [["WEEK NO:", "DATE:"], ["NO. OF STUDENTS:", "DAY:"], ["VENUE / LAB NO:", "DURATION (MINS):"]]
    for r in range(3):
        admin_table.cell(r, 0).paragraphs[0].add_run(labels[r][0]).bold = True
        admin_table.cell(r, 2).paragraphs[0].add_run(labels[r][1]).bold = True
    doc.add_paragraph()

    # 2. Resources Table
    p_res = doc.add_paragraph()
    p_res.add_run("RESOURCES & MATERIALS").bold = True
    res_table = doc.add_table(rows=1, cols=1)
    res_table.style = 'Table Grid'
    res_cell_p = res_table.cell(0, 0).paragraphs[0]
    res_cell_p.paragraph_format.line_spacing = 1.0
    res_cell_p.add_run("Smart board, Chromebook, Writing table, Projector, Screen share with laptop")
    doc.add_paragraph()

    # 3. Content Parsing & Table Boxing with Asterisk Filtering
    sections = text.split('SECTION:')
    for section in sections:
        if not section.strip(): 
            continue
            
        lines = section.strip().split('\n')
        title = lines[0].strip().upper().replace("**", "")  # Enforce FULL CAPITAL LETTERS for titles
        body_content = "\n".join(lines[1:])

        if "PEDATI" in title or "FLOW GRID" in title:
            p_sec = doc.add_paragraph()
            p_sec.add_run("P.E.D.A.T.I FLOW BREAKDOWN").bold = True
            
            # Split into individual structural blocks
            blocks = body_content.split("BLOCK_START:")
            for block in blocks:
                if not block.strip(): 
                    continue
                
                # Isolate the content before BLOCK_END
                actual_block_content = block.split("BLOCK_END")[0].strip()
                block_lines = actual_block_content.split('\n')
                
                heading_title = block_lines[0].strip().upper().replace("**", "")
                
                # State accumulation logic tracking markers
                lecturer_content = []
                students_content = []
                current_target = None
                
                for bline in block_lines[1:]:
                    cleaned_line = bline.strip()
                    if cleaned_line.upper().startswith("LECTURER:"):
                        current_target = "LECTURER"
                        continue
                    elif cleaned_line.upper().startswith("STUDENTS:"):
                        current_target = "STUDENTS"
                        continue
                    
                    if current_target == "LECTURER":
                        lecturer_content.append(cleaned_line)
                    elif current_target == "STUDENTS":
                        students_content.append(cleaned_line)

                # Assemble extracted text safely
                lecturer_text = "\n".join([l for l in lecturer_content if l])
                students_text = "\n".join([l for l in students_content if l])
                
                bp = doc.add_paragraph()
                brun = bp.add_run(heading_title)
                brun.bold = True
                
                # Draw grid layout structure box
                table = doc.add_table(rows=2, cols=2)
                table.style = 'Table Grid'
                
                for row in table.rows:
                    row.cells[0].width = Inches(3.25)
                    row.cells[1].width = Inches(3.25)
                
                hdr_cells = table.rows[0].cells
                hdr_cells[0].paragraphs[0].add_run("LECTURER").bold = True
                hdr_cells[1].paragraphs[0].add_run("STUDENTS").bold = True
                
                c0_p = table.rows[1].cells[0].paragraphs[0]
                c0_p.paragraph_format.line_spacing = 1.0
                c1_p = table.rows[1].cells[1].paragraphs[0]
                c1_p.paragraph_format.line_spacing = 1.0
                
                c0_p.add_run(lecturer_text if lecturer_text else "No content generated.")
                c1_p.add_run(students_text if students_text else "No content generated.")
                doc.add_paragraph()
        else:
            p_sec = doc.add_paragraph()
            p_sec.add_run(title).bold = True
            
            table = doc.add_table(rows=1, cols=1)
            table.style = 'Table Grid'
            cell_p = table.cell(0, 0).paragraphs[0]
            cell_p.paragraph_format.line_spacing = 1.0
            
            # REPAIRED: Reads body content directly without splitting lines to avoid text dropping
            cleaned_body = body_content.strip().replace("**", "")
            cell_p.add_run(cleaned_body if cleaned_body else "No content generated.")
            doc.add_paragraph()

    # 4. HOD Approval Page
    doc.add_page_break()
    p_hod = doc.add_paragraph()
    p_hod.add_run("HOD APPROVAL & REMARKS").bold = True
    
    hod_table = doc.add_table(rows=3, cols=2)
    hod_table.style = 'Table Grid'
    hod_table.cell(0, 0).paragraphs[0].add_run("REMARK").bold = True
    hod_table.cell(0, 1).paragraphs[0].add_run("SIGNATURE / STAMP").bold = True
    hod_table.rows[1].height = Pt(40)
    hod_table.cell(2, 0).paragraphs[0].add_run("DATE:").bold = True
    hod_table.cell(2, 1).paragraphs[0].add_run("NAME:").bold = True

    # Adjust spacing across all generated grid cells
    for row in admin_table.rows:
        for cell in row.cells: cell.paragraphs[0].paragraph_format.line_spacing = 1.0
    for row in hod_table.rows:
        for cell in row.cells: cell.paragraphs[0].paragraph_format.line_spacing = 1.0

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


# --- 4. MAIN GUI INTERFACE ---
st.write("---")
u_topic = st.text_input("LESSON TOPIC:")
u_syllabus = st.text_input("SYLLABUS CODE:")
u_extra = st.text_area("SPECIFIC CONTEXT / KEYWORDS (OPTIONAL):")

if st.button("🚀 GENERATE MASTER LESSON PLAN", type="primary"):
    if not user_api_key:
        st.error("❌ KEY CONFIGURATION ERROR! PLEASE INPUT YOUR GOOGLE GEMINI API KEY AT THE TOP OF THE PAGE FIRST.")
    elif not u_topic or not u_syllabus:
        st.error("❌ PLEASE PROVIDE BOTH A LESSON TOPIC AND A SYLLABUS CODE.")
    else:
        with st.spinner("AI IS BUILDING YOUR MASTER PLAN..."):
            result = generate_pedati_plan(u_topic, u_syllabus, u_extra, user_api_key, selected_model_name)
            st.session_state['pedati_out'] = result

if 'pedati_out' in st.session_state:
    st.divider()
    st.subheader("👁️ AI PREVIEW")
    st.text_area("GENERATED CONTENT CONTENT PREVIEW", st.session_state['pedati_out'], height=350)
    
    doc_file = create_word_export(u_topic, u_syllabus, st.session_state['pedati_out'])
    st.download_button(
        label="📥 DOWNLOAD WORD (.DOCX)", 
        data=doc_file, 
        file_name=f"LP_Merged_{u_topic.upper().replace(' ', '_')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# --- FOOTER SECTION ---
st.markdown("---") 
st.markdown(
    """
    <div style='text-align: center; color: grey; font-size: 0.8em;'>
        <p><b>MERGED SMART LESSON PLAN AI-GENERATOR V1.0</b></p>
        <p>DEVELOPED & CONCEPTUALIZED BY: <b>[HAJAH NURUL HAZIQAH @ HJH HARTINI HJ NORDIN]</b></p>
        <p>© 2026 BSC(HONORS) IN COMPUTER SCIENCE, UNIVERSITY OF STRATHCLYDE</p>
    </div>
    """,
    unsafe_allow_html=True
)
